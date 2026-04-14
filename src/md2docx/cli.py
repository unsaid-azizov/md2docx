#!/usr/bin/env python3
"""
md2docx — convert a Markdown file (with Mermaid diagrams) to .docx

Usage:
    md2docx <input.md> [output_dir] [options]

    input.md    — path to the source Markdown file
    output_dir  — where to write results (default: <input_stem>_output
                  next to the input file)

Output layout:
    <output_dir>/
        <stem>.docx        — final document
        diagrams/          — rendered Mermaid PNG files

Pipeline:
    1. mmdc   — render Mermaid code blocks to PNG (theme neutral)
    2. merge  — move standalone *Рисунок N — ...* captions into image alt text
    3. pandoc — convert processed markdown to docx
    4. docx   — font/size applied, images downscaled only, table borders
"""

import argparse
import json
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

PUPPETEER_CFG_NAME = "_puppeteer.json"

# Defaults
DEFAULT_MERMAID_WIDTH = 2400
DEFAULT_FONT = "Times New Roman"
DEFAULT_FONT_SIZE = 12
DEFAULT_MAX_WIDTH_CM = 14.0
DEFAULT_MAX_HEIGHT_CM = 20.0


# ---------------------------------------------------------------------------
# Step 1: render Mermaid diagrams
# ---------------------------------------------------------------------------

def run_mmdc(src_md: Path, rendered_md: Path, diagrams_dir: Path,
             width: int = DEFAULT_MERMAID_WIDTH) -> None:
    print(f"▶ Rendering Mermaid diagrams (theme neutral, width {width}px)...")
    diagrams_dir.mkdir(parents=True, exist_ok=True)

    # Write puppeteer config for headless Chromium in Docker (no-sandbox)
    cfg_path = diagrams_dir / PUPPETEER_CFG_NAME
    cfg_path.write_text(json.dumps({"args": ["--no-sandbox", "--disable-setuid-sandbox"]}))

    result = subprocess.run(
        [
            "mmdc",
            "-i", str(src_md),
            "-o", str(rendered_md),
            "--outputFormat", "png",
            "--theme", "neutral",
            "--width", str(width),
            "--backgroundColor", "white",
            "--puppeteerConfigFile", str(cfg_path),
        ],
        capture_output=True,
        text=True,
    )
    cfg_path.unlink(missing_ok=True)

    if result.stderr.strip():
        print(f"mmdc: {result.stderr.strip()}", file=sys.stderr)
    if not rendered_md.exists():
        print("ERROR: mmdc produced no output", file=sys.stderr)
        sys.exit(1)


# ---------------------------------------------------------------------------
# Step 2: merge captions
# ---------------------------------------------------------------------------

def process_markdown(text: str) -> str:
    """
    1. Strip the default "diagram" alt text that mmdc injects into every
       rendered Mermaid image — without this pandoc prints "diagram" as a
       figure caption for every unlabelled diagram.
    2. Move standalone *Рисунок N — caption* / *Figure N — caption* lines
       that precede an image into the image's alt text so pandoc renders
       them as proper figure captions.
    """
    # Remove mmdc default alt text
    text = re.sub(r"!\[diagram\]\(", "![](", text)

    # Merge explicit captions into the following image's alt text
    text = re.sub(
        r"\*([^\*\n]+)\*[ \t]*\n+[ \t]*!\[[^\]]*\]\(([^)]+)\)",
        lambda m: f"![{m.group(1).strip()}]({m.group(2)})",
        text,
    )
    return text


# ---------------------------------------------------------------------------
# Step 3: pandoc conversion
# ---------------------------------------------------------------------------

def run_pandoc(processed_md: Path, output_docx: Path,
               src_dir: Path, diagrams_dir: Path) -> None:
    print("▶ Converting to docx (pandoc)...")
    resource_path = f"{diagrams_dir}:{src_dir}"
    result = subprocess.run(
        [
            "pandoc",
            str(processed_md),
            "--from", "markdown",
            "--to", "docx",
            "--output", str(output_docx),
            "--resource-path", resource_path,
            "--standalone",
        ],
        capture_output=True,
        text=True,
        cwd=str(src_dir),
    )
    if result.stderr.strip():
        print(result.stderr.strip())
    if result.returncode != 0:
        sys.exit(1)


# ---------------------------------------------------------------------------
# Step 4: python-docx post-processing
# ---------------------------------------------------------------------------

def _set_font(para, name: str, size_pt: int) -> None:
    for run in para.runs:
        run.font.name = name
        run.font.size = Pt(size_pt)
        rPr = run._element.get_or_add_rPr()
        rf = rPr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rPr.insert(0, rf)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rf.set(qn(attr), name)

    pPr = para._element.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.append(rPr)
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.insert(0, rf)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(attr), name)
    half = str(size_pt * 2)
    for tag in ("w:sz", "w:szCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:val"), half)


def _has_drawing(para) -> bool:
    return bool(para._element.findall(".//" + qn("w:drawing")))


def _cap_images(doc: Document, max_width_cm: float, max_height_cm: float) -> None:
    """Downscale images to fit within max dimensions. Never enlarge."""
    max_w = int(max_width_cm / 2.54 * 914400)
    max_h = int(max_height_cm / 2.54 * 914400)
    for shape in doc.inline_shapes:
        w, h = shape.width, shape.height
        if not w or not h:
            continue
        scale = min(1.0, max_w / w, max_h / h)
        if scale < 1.0:
            shape.width = int(w * scale)
            shape.height = int(h * scale)


def _add_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    ex = tblPr.find(qn("w:tblBorders"))
    if ex is not None:
        tblPr.remove(ex)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def postprocess(output_docx: Path, font: str, font_size: int,
                max_width_cm: float, max_height_cm: float) -> None:
    print(f"▶ Post-processing: {font} {font_size}pt, images ≤{max_width_cm}×{max_height_cm}cm, table borders...")
    doc = Document(str(output_docx))

    try:
        doc.styles["Normal"].font.name = font
        doc.styles["Normal"].font.size = Pt(font_size)
    except Exception:
        pass

    prev_image = False
    for para in doc.paragraphs:
        _set_font(para, font, font_size)
        has_image = _has_drawing(para)
        if has_image or prev_image:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if para.style and "caption" in para.style.name.lower():
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        prev_image = has_image

    _cap_images(doc, max_width_cm, max_height_cm)

    for table in doc.tables:
        _add_borders(table)
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _set_font(para, font, font_size)

    doc.save(str(output_docx))


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

DOCTOR_TEXT = f"""
md2docx — Markdown + Mermaid → .docx

USAGE
  md2docx <input.md> [output_dir] [options]

  input.md        path to source Markdown file
  output_dir      where to write results
                  (default: <stem>_output/ next to the input file)

OPTIONS
  --width PX      Mermaid render width in px      (default: {DEFAULT_MERMAID_WIDTH})
  --font NAME     body font name                  (default: {DEFAULT_FONT})
  --font-size PT  body font size in points        (default: {DEFAULT_FONT_SIZE})
  --max-width CM  max image display width in cm   (default: {DEFAULT_MAX_WIDTH_CM})
  --max-height CM max image display height in cm  (default: {DEFAULT_MAX_HEIGHT_CM})

EXAMPLES
  md2docx thesis.md
  md2docx report.md /tmp/out --font "Arial" --font-size 11
  md2docx arch.md --width 3600 --max-width 16

OUTPUT
  <output_dir>/
    <stem>.docx      — final document
    diagrams/        — rendered Mermaid PNG files

SYSTEM DEPENDENCIES  (must be installed separately)
  pandoc      brew install pandoc
  mmdc        npm install -g @mermaid-js/mermaid-cli

UPDATE THIS TOOL
  uv tool install git+https://github.com/unsaid-azizov/md2docx.git --reinstall

UNINSTALL
  uv tool uninstall md2docx
"""


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convert Markdown + Mermaid to .docx",
        add_help=True,
    )
    parser.add_argument("input", nargs="?", help="Path to source .md file")
    parser.add_argument(
        "output_dir",
        nargs="?",
        default=None,
        help="Output directory (default: <input_stem>_output next to input file)",
    )
    parser.add_argument(
        "--width",
        type=int,
        default=DEFAULT_MERMAID_WIDTH,
        metavar="PX",
        help=f"Mermaid render width in pixels (default: {DEFAULT_MERMAID_WIDTH})",
    )
    parser.add_argument(
        "--font",
        default=DEFAULT_FONT,
        metavar="NAME",
        help=f"Body font name (default: {DEFAULT_FONT!r})",
    )
    parser.add_argument(
        "--font-size",
        type=int,
        default=DEFAULT_FONT_SIZE,
        metavar="PT",
        dest="font_size",
        help=f"Body font size in points (default: {DEFAULT_FONT_SIZE})",
    )
    parser.add_argument(
        "--max-width",
        type=float,
        default=DEFAULT_MAX_WIDTH_CM,
        metavar="CM",
        dest="max_width",
        help=f"Max image display width in cm (default: {DEFAULT_MAX_WIDTH_CM})",
    )
    parser.add_argument(
        "--max-height",
        type=float,
        default=DEFAULT_MAX_HEIGHT_CM,
        metavar="CM",
        dest="max_height",
        help=f"Max image display height in cm (default: {DEFAULT_MAX_HEIGHT_CM})",
    )
    parser.add_argument(
        "--info", "--doctor",
        action="store_true",
        help="Show usage, dependencies, and update instructions",
    )
    args = parser.parse_args()

    if args.info or not args.input:
        print(DOCTOR_TEXT)
        sys.exit(0)

    src_md = Path(args.input).resolve()
    if not src_md.exists():
        print(f"ERROR: file not found: {src_md}", file=sys.stderr)
        sys.exit(1)

    src_dir = src_md.parent
    stem = src_md.stem

    if args.output_dir:
        out_dir = Path(args.output_dir).resolve()
    else:
        out_dir = src_dir / f"{stem}_output"

    diagrams_dir = out_dir / "diagrams"
    rendered_md = diagrams_dir / "_rendered.md"
    processed_md = diagrams_dir / "_processed.md"
    output_docx = out_dir / f"{stem}.docx"

    out_dir.mkdir(parents=True, exist_ok=True)

    run_mmdc(src_md, rendered_md, diagrams_dir, width=args.width)

    print("▶ Processing captions...")
    processed = process_markdown(rendered_md.read_text(encoding="utf-8"))
    processed_md.write_text(processed, encoding="utf-8")

    run_pandoc(processed_md, output_docx, src_dir, diagrams_dir)
    postprocess(output_docx,
                font=args.font,
                font_size=args.font_size,
                max_width_cm=args.max_width,
                max_height_cm=args.max_height)

    # Clean up temp markdown files; keep PNGs in diagrams/
    rendered_md.unlink(missing_ok=True)
    processed_md.unlink(missing_ok=True)

    print(f"✓ Done: {output_docx}")
    print(f"  Diagrams: {diagrams_dir}/")


if __name__ == "__main__":
    main()
